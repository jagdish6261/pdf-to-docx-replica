
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="PDF to DOCX Replica", layout="centered")

st.title("📄 PDF to MS Word Replica Tool")
st.write("This tool recreates a structured MS Word document from a predefined PDF layout.")

if st.button("Generate Word Document"):
    doc = Document()

    def center_bold(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def bold(text):
        run = doc.add_paragraph().add_run(text)
        run.bold = True

    center_bold("FORM ‘A’")
    center_bold("MEDIATION APPLICATION FORM")
    center_bold("[REFER RULE 3(1)]")
    center_bold("Mumbai District Legal Services Authority")
    center_bold("City Civil Court, Mumbai")

    bold("DETAILS OF PARTIES:")

    bold("1. Name of Applicant")
    doc.add_paragraph("{{client_name}}")

    bold("Address and Contact Details of Applicant")
    doc.add_paragraph("REGISTERED ADDRESS:\n{{branch_address}}")
    doc.add_paragraph("CORRESPONDENCE BRANCH ADDRESS:\n{{branch_address}}")
    doc.add_paragraph("Telephone No.: {{mobile}}")
    doc.add_paragraph("Email ID: info@kslegal.co.in")

    bold("2. Name, Address and Contact details of Opposite Party")
    doc.add_paragraph("Name: {{customer_name}}")

    doc.add_paragraph("REGISTERED ADDRESS:\n________________")
    doc.add_paragraph("CORRESPONDENCE ADDRESS:\n________________")

    bold("DETAILS OF DISPUTE:")
    bold("THE COMMERCIAL COURTS (PRE-INSTITUTION MEDIATION AND SETTLEMENT) RULES, 2018")

    doc.add_paragraph(
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
    )

    file_path = "Mediation_Application_Form.docx"
    doc.save(file_path)

    with open(file_path, "rb") as file:
        st.download_button(
            label="⬇ Download MS Word File",
            data=file,
            file_name=file_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
